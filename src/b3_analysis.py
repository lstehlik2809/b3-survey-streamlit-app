from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import networkx as nx
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from unidecode import unidecode


class B3InputError(ValueError):
    """Raised when an uploaded workbook does not match the B3 input format."""


@dataclass(frozen=True)
class B3ReportResult:
    report_path: Path
    positive_chart_path: Path
    negative_chart_path: Path
    positive_table: pd.DataFrame
    negative_table: pd.DataFrame
    positive_network: dict[str, Any]
    negative_network: dict[str, Any]
    stats: dict[str, int]
    layout_engine: str


REQUIRED_SHEETS = {"nodes", "edges"}
REQUIRED_NODE_COLUMNS = {
    "nodeId",
    "gender",
    "presentWithConsent",
    "positivePoints",
    "negativePoints",
}
REQUIRED_EDGE_COLUMNS = {
    "source",
    "target",
    "relationshipValence",
    "relationshipStrength",
}

RELATIONSHIP_CONFIG = {
    "pos": {
        "relationshipValence": "positive",
        "pointsType": "positivePoints",
        "relationshipType": "Kladné vztahy",
        "outputName": "kladneVztahy.png",
    },
    "neg": {
        "relationshipValence": "negative",
        "pointsType": "negativePoints",
        "relationshipType": "Záporné vztahy",
        "outputName": "zaporneVztahy.png",
    },
}

CHART_FIGSIZE = (9, 7)
REPORT_CHART_WIDTH = Inches(6.5)
GRAPHVIZ_LAYOUT_PROGRAM = "sfdp"
GRAPHVIZ_LAYOUT_ARGS = "-Goverlap=prism -Gsep=+18 -GK=0.8"


def camel_case(value: str) -> str:
    output = "".join(part for part in value.title() if part.isalnum())
    return output[0].lower() + output[1:] if output else output


def generate_b3_report(workbook_path: str | Path, output_dir: str | Path) -> B3ReportResult:
    """Generate a B3 DOCX report from a workbook matching B3Inputs.xlsx."""
    workbook_path = Path(workbook_path)
    output_dir = Path(output_dir)
    charts_dir = output_dir / "outputCharts"
    report_dir = output_dir / "outputReport"
    charts_dir.mkdir(parents=True, exist_ok=True)
    report_dir.mkdir(parents=True, exist_ok=True)

    nodes, edges = _read_b3_workbook(workbook_path)
    _validate_node_values(nodes)
    _validate_edges_reference_nodes(nodes, edges)

    overview_parts: list[pd.DataFrame] = []
    chart_paths: dict[str, Path] = {}
    networks: dict[str, dict[str, Any]] = {}
    layout_engine = f"graphviz/{GRAPHVIZ_LAYOUT_PROGRAM}"

    for key, config in RELATIONSHIP_CONFIG.items():
        filtered_edges = edges[
            edges["relationshipValence"] == config["relationshipValence"]
        ].copy()

        all_nodes_points = _score_nodes(nodes, filtered_edges, config["pointsType"])
        node_attrs = nodes[["nodeId", "gender", "presentWithConsent"]].merge(
            all_nodes_points, on="nodeId", how="left"
        )

        graph = _build_relationship_graph(node_attrs, filtered_edges)
        chart_path = charts_dir / config["outputName"]
        positions, used_engine = _get_layout_positions(graph)
        _draw_relationship_chart(graph, positions, config["relationshipType"], chart_path)
        if used_engine != f"graphviz/{GRAPHVIZ_LAYOUT_PROGRAM}":
            layout_engine = used_engine
        chart_paths[key] = chart_path
        networks[key] = _network_to_interactive_data(
            graph, positions, config["relationshipType"]
        )

        table = node_attrs[["nodeId", "pointsOverall"]].sort_values(
            by="pointsOverall", axis=0, ascending=False
        )
        table = table.rename(columns={"nodeId": "ID žáka", "pointsOverall": "Počet bodů"})
        table["valence"] = config["relationshipValence"]
        overview_parts.append(table)

    tab_overview = pd.concat(overview_parts, ignore_index=True)
    positive_table = (
        tab_overview[tab_overview["valence"] == "positive"]
        .drop("valence", axis=1)
        .rename(columns={"Počet bodů": "Počet kladných bodů"})
    )
    negative_table = (
        tab_overview[tab_overview["valence"] == "negative"]
        .drop("valence", axis=1)
        .rename(columns={"Počet bodů": "Počet záporných bodů"})
    )

    stats = {
        "total_students": int(len(nodes)),
        "male_students": int(len(nodes[nodes["gender"] == "male"])),
        "female_students": int(len(nodes[nodes["gender"] == "female"])),
        "not_present_or_without_consent": int(
            len(nodes[nodes["presentWithConsent"] == "no"])
        ),
    }

    report_path = report_dir / "B3_Report.docx"
    _write_docx_report(
        nodes=nodes,
        positive_chart_path=chart_paths["pos"],
        negative_chart_path=chart_paths["neg"],
        positive_table=positive_table,
        negative_table=negative_table,
        report_path=report_path,
    )

    return B3ReportResult(
        report_path=report_path,
        positive_chart_path=chart_paths["pos"],
        negative_chart_path=chart_paths["neg"],
        positive_table=positive_table,
        negative_table=negative_table,
        positive_network=networks["pos"],
        negative_network=networks["neg"],
        stats=stats,
        layout_engine=layout_engine,
    )


def _read_b3_workbook(workbook_path: Path) -> tuple[pd.DataFrame, pd.DataFrame]:
    if not workbook_path.exists():
        raise B3InputError(f"Input file does not exist: {workbook_path}")

    try:
        excel_file = pd.ExcelFile(workbook_path)
        sheet_names = excel_file.sheet_names
    except Exception as exc:
        raise B3InputError("The uploaded file could not be opened as an Excel workbook.") from exc

    try:
        missing_sheets = REQUIRED_SHEETS.difference(sheet_names)
        if missing_sheets:
            missing = ", ".join(sorted(missing_sheets))
            raise B3InputError(f"The workbook is missing required sheet(s): {missing}.")

        nodes = pd.read_excel(excel_file, sheet_name="nodes", header=0)
        edges = pd.read_excel(excel_file, sheet_name="edges", header=0)
    finally:
        excel_file.close()

    _require_columns("nodes", nodes, REQUIRED_NODE_COLUMNS)
    _require_columns("edges", edges, REQUIRED_EDGE_COLUMNS)

    nodes = nodes.copy()
    edges = edges.copy()
    nodes["positivePoints"] = pd.to_numeric(nodes["positivePoints"], errors="coerce").fillna(0)
    nodes["negativePoints"] = pd.to_numeric(nodes["negativePoints"], errors="coerce").fillna(0)
    edges["relationshipStrength"] = pd.to_numeric(
        edges["relationshipStrength"], errors="coerce"
    ).fillna(0)

    if nodes["nodeId"].isna().any():
        raise B3InputError("The nodes sheet contains empty nodeId values.")
    if edges[["source", "target"]].isna().any().any():
        raise B3InputError("The edges sheet contains empty source or target values.")

    return nodes, edges


def _require_columns(sheet_name: str, frame: pd.DataFrame, required: set[str]) -> None:
    missing_columns = required.difference(frame.columns)
    if missing_columns:
        missing = ", ".join(sorted(missing_columns))
        raise B3InputError(f"The {sheet_name} sheet is missing required column(s): {missing}.")


def _validate_node_values(nodes: pd.DataFrame) -> None:
    unknown_gender = set(nodes["gender"].dropna().unique()).difference({"male", "female"})
    if unknown_gender:
        values = ", ".join(sorted(map(str, unknown_gender)))
        raise B3InputError(f"Unknown gender value(s) in nodes sheet: {values}.")

    unknown_consent = set(nodes["presentWithConsent"].dropna().unique()).difference(
        {"yes", "no"}
    )
    if unknown_consent:
        values = ", ".join(sorted(map(str, unknown_consent)))
        raise B3InputError(f"Unknown presentWithConsent value(s) in nodes sheet: {values}.")


def _validate_edges_reference_nodes(nodes: pd.DataFrame, edges: pd.DataFrame) -> None:
    node_ids = set(nodes["nodeId"])
    edge_ids = set(edges["source"]).union(edges["target"])
    missing_ids = edge_ids.difference(node_ids)
    if missing_ids:
        values = ", ".join(sorted(map(str, missing_ids)))
        raise B3InputError(f"Edge source/target value(s) are missing from nodes.nodeId: {values}.")


def _score_nodes(
    nodes: pd.DataFrame, filtered_edges: pd.DataFrame, points_type: str
) -> pd.DataFrame:
    all_nodes = pd.DataFrame({"nodeId": nodes["nodeId"]})
    edge_points = (
        filtered_edges.groupby("target", dropna=True)["relationshipStrength"].sum().reset_index()
    )
    all_nodes_points = all_nodes.merge(
        edge_points, left_on="nodeId", right_on="target", how="left"
    )
    all_nodes_points = all_nodes_points[["nodeId", "relationshipStrength"]]
    all_nodes_points["relationshipStrength"] = all_nodes_points[
        "relationshipStrength"
    ].fillna(0)

    additional_points = nodes.groupby("nodeId", dropna=True)[points_type].sum().reset_index()
    all_nodes_points = all_nodes_points.merge(additional_points, on="nodeId", how="left")
    all_nodes_points[points_type] = all_nodes_points[points_type].fillna(0)
    all_nodes_points["pointsOverall"] = (
        all_nodes_points["relationshipStrength"] + all_nodes_points[points_type]
    )
    all_nodes_points = all_nodes_points[["nodeId", "pointsOverall"]]
    all_nodes_points["pointsOverall"] = all_nodes_points["pointsOverall"].astype(int)
    return all_nodes_points


def _build_relationship_graph(
    node_attrs: pd.DataFrame, filtered_edges: pd.DataFrame
) -> nx.DiGraph:
    graph = nx.DiGraph()
    graph.add_nodes_from(node_attrs["nodeId"].to_numpy())
    graph.add_edges_from(list(zip(filtered_edges["source"], filtered_edges["target"])))
    nx.set_node_attributes(graph, node_attrs.set_index("nodeId").to_dict("index"))
    return graph


def _draw_relationship_chart(
    graph: nx.DiGraph,
    positions: dict[Any, tuple[float, float]],
    relationship_type: str,
    chart_path: Path,
) -> None:
    nodes = list(graph.nodes())
    gender_colors = {"male": "blue", "female": "yellow"}
    node_colors = [gender_colors[graph.nodes[node]["gender"]] for node in nodes]

    edge_node_colors = {"yes": "grey", "no": "red"}
    outline_colors = [
        edge_node_colors[graph.nodes[node]["presentWithConsent"]] for node in nodes
    ]

    score_labels = nx.get_node_attributes(graph, "pointsOverall")

    plt.figure(figsize=CHART_FIGSIZE)
    plt.title(relationship_type)
    nx.draw_networkx_nodes(
        graph,
        pos=positions,
        node_color=node_colors,
        alpha=0.5,
        node_shape="o",
        edgecolors=outline_colors,
    )
    nx.draw_networkx_edges(graph, pos=positions, edge_color="grey", width=0.5)
    nx.draw_networkx_labels(graph, pos=positions, font_size=7, verticalalignment="top")
    nx.draw_networkx_labels(
        graph,
        pos=positions,
        labels=score_labels,
        font_size=5,
        verticalalignment="bottom",
    )
    plt.gca().set_aspect("equal", adjustable="box")
    plt.axis("off")
    plt.savefig(chart_path, dpi=500)
    plt.close()


def _get_layout_positions(graph: nx.DiGraph) -> tuple[dict[Any, tuple[float, float]], str]:
    try:
        raw_positions = nx.nx_agraph.graphviz_layout(
            graph, prog=GRAPHVIZ_LAYOUT_PROGRAM, args=GRAPHVIZ_LAYOUT_ARGS
        )
        return (
            _map_layout_keys_to_graph_nodes(raw_positions, graph),
            f"graphviz/{GRAPHVIZ_LAYOUT_PROGRAM}",
        )
    except Exception:
        return _component_kamada_kawai_layout(graph), "networkx/kamada_kawai_layout"


def _component_kamada_kawai_layout(graph: nx.DiGraph) -> dict[Any, tuple[float, float]]:
    undirected_graph = graph.to_undirected()
    components = sorted(nx.connected_components(undirected_graph), key=len, reverse=True)
    positions: dict[Any, tuple[float, float]] = {}

    cursor_x = 0.0
    gap = 3.0

    for component_nodes in components:
        subgraph = undirected_graph.subgraph(component_nodes)
        component_size = len(component_nodes)

        if component_size == 1:
            local_positions = {next(iter(component_nodes)): (0.0, 0.0)}
            component_scale = 1.0
        else:
            component_scale = max(1.4, component_size / 4)
            local_positions = nx.kamada_kawai_layout(subgraph, scale=component_scale)

        min_x = min(position[0] for position in local_positions.values())
        max_x = max(position[0] for position in local_positions.values())
        min_y = min(position[1] for position in local_positions.values())
        max_y = max(position[1] for position in local_positions.values())
        width = max(max_x - min_x, 1.0)
        center_y = (min_y + max_y) / 2

        for node, position in local_positions.items():
            positions[node] = (position[0] - min_x + cursor_x, position[1] - center_y)

        cursor_x += width + gap

    return positions


def _map_layout_keys_to_graph_nodes(
    raw_positions: dict[Any, tuple[float, float]], graph: nx.DiGraph
) -> dict[Any, tuple[float, float]]:
    lookup = {node: node for node in graph.nodes()}
    lookup.update({str(node): node for node in graph.nodes()})
    return {lookup.get(key, key): value for key, value in raw_positions.items()}


def _network_to_interactive_data(
    graph: nx.DiGraph,
    positions: dict[Any, tuple[float, float]],
    relationship_type: str,
) -> dict[str, Any]:
    return {
        "title": relationship_type,
        "nodes": [
            {
                "id": str(node),
                "label": str(node),
                "x": float(positions[node][0]),
                "y": float(positions[node][1]),
                "gender": str(graph.nodes[node]["gender"]),
                "presentWithConsent": str(graph.nodes[node]["presentWithConsent"]),
                "pointsOverall": int(graph.nodes[node]["pointsOverall"]),
            }
            for node in graph.nodes()
        ],
        "edges": [
            {
                "source": str(source),
                "target": str(target),
            }
            for source, target in graph.edges()
        ],
    }


def _write_docx_report(
    *,
    nodes: pd.DataFrame,
    positive_chart_path: Path,
    negative_chart_path: Path,
    positive_table: pd.DataFrame,
    negative_table: pd.DataFrame,
    report_path: Path,
) -> None:
    doc = Document()

    doc.add_heading("Grafický výstup z dotazníku B-3", 0)

    doc.add_paragraph(
        "Dva níže uvedené grafy zobrazují vztahy mezi členy třídního kolektivu "
        "na základě jejich pozitivních a negativních voleb týkajících se toho, "
        "koho ze svých spolužáků považují za své přátele."
    )

    stats = doc.add_paragraph()
    stats.add_run("Celkový počet žáků: ").bold = True
    stats.add_run(str(len(nodes)))
    stats.add_run("\nPočet chlapců a dívek: ").bold = True
    stats.add_run(
        str(len(nodes[nodes["gender"] == "male"]))
        + " / "
        + str(len(nodes[nodes["gender"] == "female"]))
    )
    stats.add_run("\nPočet nepřítomných žáků nebo bez souhlasu: ").bold = True
    stats.add_run(str(len(nodes[nodes["presentWithConsent"] == "no"])))

    intro = doc.add_paragraph()
    intro.add_run("Vysvětlivky ke grafům:").bold = True
    intro.add_run(
        "\n "
        + chr(9679)
        + " Směr šipek reprezentuje, kdo koho označil ve své pozitivní nebo negativní volbě."
    )
    intro.add_run(
        "\n "
        + chr(9679)
        + " Barva uzlů odpovídá pohlaví žáků - žlutá barva reprezentuje dívky a fialová barva chlapce."
    )
    intro.add_run(
        "\n "
        + chr(9679)
        + " Uzly s červeným okrajem odpovídají žákům, kteří nebyli v době sběru dat přítomni nebo u nich není k dispozici souhlas se zpracováním jejich odpovědí."
    )
    intro.add_run(
        "\n "
        + chr(9679)
        + " Spodní číslo v uzlu odpovídá kódovému označení žáka."
    )
    intro.add_run(
        "\n "
        + chr(9679)
        + " Horní číslo v uzlu odpovídá součtu vážených odpovědí na otázky č. 1., resp. 2, a 6."
    )
    intro.add_run(
        "\n "
        + chr(9679)
        + " Vzdálenosti mezi uzly jsou zvoleny tak, aby co nejlépe odrážely vnitřní strukturu třídního kolektivu, tzn. tak, aby žáci s podobnými a vzájemnýmmi volbami byli v grafu blízko sebe."
    )

    doc.add_picture(str(positive_chart_path), width=REPORT_CHART_WIDTH)
    doc.add_picture(str(negative_chart_path), width=REPORT_CHART_WIDTH)

    tab_pos_intro = doc.add_paragraph()
    tab_pos_intro.add_run(
        "Žáci sestupně seřazení podle počtu obdržených kladných bodů"
    ).bold = True
    _add_dataframe_table(doc, positive_table)

    br = doc.add_paragraph()
    run = br.add_run()
    run.add_break()

    tab_neg_intro = doc.add_paragraph()
    tab_neg_intro.add_run(
        "Žáci sestupně seřazení podle počtu obdržených záporných bodů"
    ).bold = True
    _add_dataframe_table(doc, negative_table)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    doc.save(report_path)


def _add_dataframe_table(doc: Document, frame: pd.DataFrame) -> None:
    table = doc.add_table(frame.shape[0] + 1, frame.shape[1])
    table.style = "TableGrid"

    for column_index in range(frame.shape[-1]):
        table.cell(0, column_index).text = frame.columns[column_index]

    for row_index in range(frame.shape[0]):
        for column_index in range(frame.shape[-1]):
            table.cell(row_index + 1, column_index).text = str(
                frame.values[row_index, column_index]
            )


def generate_default_report() -> B3ReportResult:
    """Compatibility helper for the original local sample workbook."""
    project_root = Path(__file__).resolve().parents[1]
    return generate_b3_report(project_root / "B3Inputs.xlsx", project_root)
